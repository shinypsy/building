import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    """표 셀 배경색 채우기"""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=90, bottom=90, left=140, right=140):
    """표 셀 안쪽 여백(Padding) 설정 (dxa 단위: 1pt = 20 dxa)"""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_borders(cell, **kwargs):
    """
    셀 개별 테두리 설정 (left, top, right, bottom)
    """
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    
    for edge in ('top', 'left', 'bottom', 'right'):
        edge_data = kwargs.get(edge)
        if edge_data:
            b = OxmlElement(f'w:{edge}')
            b.set(qn('w:val'), edge_data.get('val', 'single'))
            b.set(qn('w:sz'), str(edge_data.get('sz', 4)))
            b.set(qn('w:space'), '0')
            b.set(qn('w:color'), edge_data.get('color', 'auto'))
            tcBorders.append(b)
        else:
            b = OxmlElement(f'w:{edge}')
            b.set(qn('w:val'), 'none')
            tcBorders.append(b)
    tcPr.append(tcBorders)

def add_alert_box(doc, alert_type, title, content):
    """
    마크다운의 Alert 블록(IMPORTANT, TIP, NOTE)을 1x1 테이블로 모사
    1페이지 최적화를 위해 상하 여백과 글자 크기를 극도로 정밀 제어
    """
    if alert_type == 'IMPORTANT':
        border_color = 'DC3545'  # 빨강
        bg_color = 'FDF2F3'
    elif alert_type == 'TIP':
        border_color = '28A745'  # 초록
        bg_color = 'F4F9F4'
    else:  # NOTE
        border_color = '007BFF'  # 파랑
        bg_color = 'F0F8FF'
        
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    cell = table.cell(0, 0)
    cell.width = Inches(7.1)  # 여백 0.7in 대응 가로 너비 확장
    
    set_cell_background(cell, bg_color)
    set_cell_margins(cell, top=100, bottom=100, left=180, right=140)
    set_cell_borders(cell, left={'val': 'single', 'sz': 36, 'color': border_color})
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.25
    
    run_title = p.add_run(f"★ {title}\n")
    run_title.font.name = '맑은 고딕'
    run_title.font.size = Pt(9.5)
    run_title.bold = True
    run_title.font.color.rgb = RGBColor(51, 51, 51)
    
    run_content = p.add_run(content)
    run_content.font.name = '맑은 고딕'
    run_content.font.size = Pt(8.5)
    run_content.font.color.rgb = RGBColor(85, 85, 85)

    # 1페이지 조절을 위해 AlertBox 단락 후 간격을 6pt로 제어
    p_spacer = doc.add_paragraph()
    p_spacer.paragraph_format.space_before = Pt(0)
    p_spacer.paragraph_format.space_after = Pt(4)

def main():
    doc = docx.Document()
    
    # 상하좌우 여백 0.7인치로 최소화 (1페이지 맞춤 규격)
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    # 기본 글꼴 스타일 정의 (맑은 고딕)
    style = doc.styles['Normal']
    font = style.font
    font.name = '맑은 고딕'
    font.size = Pt(9.5)
    font.color.rgb = RGBColor(51, 51, 51)
    
    # 1. 문서 제목 (16pt, Deep Navy)
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(4)
    title_p.paragraph_format.space_after = Pt(10)
    
    run = title_p.add_run("[기안서] 사내 협업 메신저 및 기업용 메일(하이웍스) 도입의 건")
    run.font.size = Pt(16)
    run.bold = True
    run.font.color.rgb = RGBColor(30, 58, 138)  # Deep Navy
    
    # 2. 결재 테이블 (1x4 구조에 2개 행 구축)
    table_meta = doc.add_table(rows=2, cols=4)
    table_meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_meta.autofit = False
    
    widths = [Inches(1.0), Inches(2.55), Inches(1.0), Inches(2.55)]
    
    row_1 = table_meta.rows[0].cells
    row_1[0].text = "기안부서"
    row_1[1].text = "관리부 / 개발기획팀"
    row_1[2].text = "기안자"
    row_1[3].text = "Jay"
    
    row_2 = table_meta.rows[1].cells
    row_2[0].text = "기안일자"
    row_2[1].text = "2026년 6월 1일"
    row_2[2].text = "결재상태"
    row_2[3].text = "대기 (결재 대기중)"
    
    for r_idx, row in enumerate(table_meta.rows):
        for c_idx, cell in enumerate(row.cells):
            cell.width = widths[c_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
            
            set_cell_borders(cell, 
                             top={'val': 'single', 'sz': 4, 'color': 'D3D3D3'},
                             bottom={'val': 'single', 'sz': 4, 'color': 'D3D3D3'},
                             left={'val': 'single', 'sz': 4, 'color': 'D3D3D3'},
                             right={'val': 'single', 'sz': 4, 'color': 'D3D3D3'})
            
            if c_idx % 2 == 0:
                set_cell_background(cell, 'F2F2F2')
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.name = '맑은 고딕'
                    run.bold = True
                    run.font.size = Pt(9.0)
            else:
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    run.font.name = '맑은 고딕'
                    run.font.size = Pt(9.0)

    # 3. 기안 목적 Section
    h1 = doc.add_paragraph()
    h1.paragraph_format.space_before = Pt(10)
    h1.paragraph_format.space_after = Pt(4)
    run_h1 = h1.add_run("1. 기안 목적")
    run_h1.font.size = Pt(11.5)
    run_h1.bold = True
    run_h1.font.color.rgb = RGBColor(30, 58, 138)
    
    p1 = doc.add_paragraph()
    p1.paragraph_format.space_after = Pt(8)
    p1.paragraph_format.line_spacing = 1.3
    run_p1 = p1.add_run(
        "현재 회사 공식 업무에 개인 이메일을 사용하는 환경은 정보 보안, 기업 대외 이미지 및 업무 몰입 방면에서 치명적 한계를 드러내고 있습니다. "
        "이에 전사적 협업 플랫폼이자 국내 1위 그룹웨어인 하이웍스(Hiworks)를 도입해 대외 신뢰도를 제고하고 안전한 협업 인프라를 구축하고자 본 안건을 기안합니다."
    )
    run_p1.font.size = Pt(9.5)
    
    # 4. 현황 및 문제점 Section
    h2 = doc.add_paragraph()
    h2.paragraph_format.space_before = Pt(10)
    h2.paragraph_format.space_after = Pt(4)
    run_h2 = h2.add_run("2. 현황 및 문제점")
    run_h2.font.size = Pt(11.5)
    run_h2.bold = True
    run_h2.font.color.rgb = RGBColor(30, 58, 138)
    
    problems = [
        ("보안 취약 및 데이터 유실 위험", "퇴사 시 업무 히스토리가 개인 소유로 남아 데이터 유실 및 기술 기밀 정보의 유출 리스크에 무방비 노출됩니다."),
        ("회사 이미지 제고 필요", "대외 바이어 및 파트너사와의 업무 소통 시 개인 메일 주소(@naver.com 등) 사용은 기업의 대외 신뢰도와 프로페셔널한 이미지 형성을 현저히 저해합니다."),
        ("공사(公私) 혼선으로 인한 몰입 방해", "카카오톡 등 일상 메신저를 업무에 혼용해 직원의 프라이버시가 침해되며, 주의 분산으로 업무 집중력이 저하됩니다.")
    ]
    
    for title, desc in problems:
        p_prob = doc.add_paragraph(style='List Bullet')
        p_prob.paragraph_format.space_before = Pt(0)
        p_prob.paragraph_format.space_after = Pt(3)
        p_prob.paragraph_format.line_spacing = 1.2
        r_title = p_prob.add_run(f" {title}: ")
        r_title.bold = True
        r_title.font.size = Pt(9.0)
        r_desc = p_prob.add_run(desc)
        r_desc.font.size = Pt(9.0)

    # 5. 개인 계정 vs 하이웍스 비교 장표 (★최우선 강조 영역)
    h4 = doc.add_paragraph()
    h4.paragraph_format.space_before = Pt(12)
    h4.paragraph_format.space_after = Pt(6)
    run_h4 = h4.add_run("3. 개인 계정 vs 하이웍스 핵심 비교 분석 (★도입 메리트)")
    run_h4.font.size = Pt(11.5)
    run_h4.bold = True
    run_h4.font.color.rgb = RGBColor(30, 58, 138)
    
    # 비교 장표 안내 Callout Box (연한 파란색 단락으로 강조 극대화)
    callout = doc.add_table(rows=1, cols=1)
    callout.alignment = WD_TABLE_ALIGNMENT.CENTER
    callout.autofit = False
    c_cell = callout.cell(0, 0)
    c_cell.width = Inches(7.1)
    set_cell_background(c_cell, 'EBF3FC') # 강조 파란색 배경
    set_cell_margins(c_cell, top=60, bottom=60, left=120, right=100)
    set_cell_borders(c_cell, left={'val': 'single', 'sz': 24, 'color': '1E3A8A'})
    
    cp = c_cell.paragraphs[0]
    cp.paragraph_format.space_before = Pt(2)
    cp.paragraph_format.space_after = Pt(2)
    c_run = cp.add_run("★ 중요 보고 요약: 하이웍스 메신저/이메일 도입 시 기대 효과 및 영역별 세부 혜택 비교")
    c_run.font.name = '맑은 고딕'
    c_run.bold = True
    c_run.font.size = Pt(9.0)
    c_run.font.color.rgb = RGBColor(30, 58, 138)
    
    # 간격 패딩
    pad = doc.add_paragraph()
    pad.paragraph_format.space_before = Pt(0)
    pad.paragraph_format.space_after = Pt(4)
    
    # 비교 테이블 구축 (가로 7.1인치를 가득 채움)
    compare_table = doc.add_table(rows=6, cols=4)
    compare_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    compare_table.autofit = False
    
    c_widths = [Inches(1.2), Inches(1.9), Inches(2.0), Inches(2.0)]
    
    headers_comp = ["비교 항목", "개인 계정 사용 (현행)", "하이웍스 도입 (제안)", "기대 효과 (도입 메리트)"]
    row_data = [
        ["데이터 소유권", "개인 소유 (퇴사 시 자료 유실)", "회사 소유 (중앙 아카이빙 보존)", "업무 인수인계 공백 제거 및 기업 자산 보호"],
        ["기업 신뢰도", "@naver.com 등 (프로필 신뢰 하락)", "@yourcompany.com (도메인 통일)", "대외 전문성 확립 및 기업 이미지 제고"],
        ["계정 보안 관리", "통제 불가능 (개인 관리 취약)", "중앙 통제 (비밀번호 강제, 접근제한)", "계정 탈취 및 사내 기밀 유출 위험 차단"],
        ["부서 협업 및 속도", "연락처 수동 검색 (연동 알림 없음)", "조직도 자동 연동 / 실시간 통합 알림", "소통 탐색 비용 제로화 및 신속한 의사결정"],
        ["소통 환경", "사설 카톡 혼용 (공사 경계 모호)", "전용 비즈니스 메신저 분리 구축", "직원 프라이버시 보호 및 업무 집중력 극대화"]
    ]
    
    # 헤더 입력 및 디자인 (Deep Navy Header)
    hdr_cells_comp = compare_table.rows[0].cells
    for c_idx, text in enumerate(headers_comp):
        hdr_cells_comp[c_idx].text = text
        set_cell_background(hdr_cells_comp[c_idx], '1E3A8A')  # Deep Navy Blue
        set_cell_margins(hdr_cells_comp[c_idx], top=80, bottom=80, left=80, right=80)
        hdr_cells_comp[c_idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        p = hdr_cells_comp[c_idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.name = '맑은 고딕'
            run.bold = True
            run.font.size = Pt(9.0)
            run.font.color.rgb = RGBColor(255, 255, 255) # 흰색 글씨

    # 내용 입력 및 디자인
    for r_idx, data in enumerate(row_data):
        row_cells = compare_table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(data):
            row_cells[c_idx].text = val
            set_cell_margins(row_cells[c_idx], top=70, bottom=70, left=80, right=80)
            row_cells[c_idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
            set_cell_borders(row_cells[c_idx],
                             top={'val': 'single', 'sz': 4, 'color': 'D3D3D3'},
                             bottom={'val': 'single', 'sz': 4, 'color': 'D3D3D3'},
                             left={'val': 'single', 'sz': 4, 'color': 'D3D3D3'},
                             right={'val': 'single', 'sz': 4, 'color': 'D3D3D3'})
            
            # Zebra pattern (짝수 줄에 연한 하늘색/회색 칠하기)
            if r_idx % 2 == 1:
                set_cell_background(row_cells[c_idx], 'F7F9FB')
            
            p = row_cells[c_idx].paragraphs[0]
            # 정렬 및 폰트 세팅
            if c_idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.name = '맑은 고딕'
                    run.bold = True
                    run.font.size = Pt(8.5)
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    run.font.name = '맑은 고딕'
                    run.font.size = Pt(8.5)
                    # 하이웍스 제안이나 기대 효과 내 강조 문구 굵게 처리
                    if c_idx >= 2 and ("회사 소유" in val or "도메인 통일" in val or "이미지 제고" in val or "자산 보호" in val):
                        run.bold = True

    # 비교 테이블 셀 너비 적용
    for row in compare_table.rows:
        for c_idx, cell in enumerate(row.cells):
            cell.width = c_widths[c_idx]

    spacer3 = doc.add_paragraph()
    spacer3.paragraph_format.space_before = Pt(0)
    spacer3.paragraph_format.space_after = Pt(10)

    # 6. 기대 효과 및 결론 Section
    h5 = doc.add_paragraph()
    h5.paragraph_format.space_before = Pt(10)
    h5.paragraph_format.space_after = Pt(4)
    run_h5 = h5.add_run("4. 기대 효과 및 결론")
    run_h5.font.size = Pt(11.5)
    run_h5.bold = True
    run_h5.font.color.rgb = RGBColor(30, 58, 138)
    
    p5 = doc.add_paragraph()
    p5.paragraph_format.space_after = Pt(4)
    p5.paragraph_format.line_spacing = 1.3
    run_p5 = p5.add_run(
        "하이웍스의 사내 메신저 및 기업용 메일 도입은 단순한 툴 교체가 아닌, 당사의 정보 자산을 보호하고 브랜드 신뢰도를 확립하기 위한 필수적 투자입니다. "
        "1페이지 요약 비교 장표에 기술된 즉각적인 경영 개선 시너지를 바탕으로 본 안건에 대한 긍정적인 검토 및 최종 승인을 바랍니다."
    )
    run_p5.font.size = Pt(9.5)

    doc.save("하이웍스_도입_기안서.docx")
    print("기안서 워드 파일(하이웍스_도입_기안서.docx) 생성 완료!")

if __name__ == '__main__':
    main()
