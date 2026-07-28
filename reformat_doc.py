from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import sys

sys.stdout.reconfigure(encoding='utf-8')

def reformat_doc(doc_path):
    """
    DOC 파일의 포맷팅을 정렬한다.
    - 글꼴 일관성 (맑은 고딕, 10pt)
    - 행 간격 (1.5줄)
    - 문단 여백
    - 들여쓰기
    - 표 정렬
    """
    print(f"📄 파일 열기: {doc_path}")
    doc = Document(doc_path)

    # 기본 스타일 설정
    style = doc.styles['Normal']
    style.font.name = '맑은 고딕'
    style.font.size = Pt(11)

    total_paragraphs = len(doc.paragraphs)
    print(f"📊 총 단락 수: {total_paragraphs}")

    # 모든 단락 정렬
    for idx, para in enumerate(doc.paragraphs):
        # 글꼴 설정
        for run in para.runs:
            if run.font.name is None or run.font.name == '':
                run.font.name = '맑은 고딕'
            if run.font.size is None:
                run.font.size = Pt(11)

        # 문단 간격 설정
        para.paragraph_format.line_spacing = 1.5
        para.paragraph_format.space_before = Pt(3)
        para.paragraph_format.space_after = Pt(3)

        # 정렬 (왼쪽 정렬이 아니면 조정)
        if para.style.name.startswith('Heading'):
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif para.alignment != WD_ALIGN_PARAGRAPH.LEFT:
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        if (idx + 1) % 50 == 0:
            print(f"  ✓ {idx + 1}/{total_paragraphs} 단락 처리 완료")

    # 표 정렬 (있으면)
    table_count = len(doc.tables)
    if table_count > 0:
        print(f"📊 표 정렬: {table_count}개")
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell in row.cells:
                    # 셀 내 단락 정렬
                    for para in cell.paragraphs:
                        para.paragraph_format.line_spacing = 1.2
                        para.paragraph_format.space_before = Pt(2)
                        para.paragraph_format.space_after = Pt(2)
                        for run in para.runs:
                            if run.font.name is None:
                                run.font.name = '맑은 고딕'
                            if run.font.size is None:
                                run.font.size = Pt(10)

    # 페이지 여백 설정 (1인치)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # 저장
    print(f"💾 파일 저장 중...")
    doc.save(doc_path)
    print(f"✅ 정렬 완료: {doc_path}")

if __name__ == '__main__':
    doc_file = r"d:\Dev\Project\sample\doc\정보통신설비_2026성능점검결과서_엠코지니어스타.docx"
    reformat_doc(doc_file)
