"""정보통신설비 유지보수·관리 / 성능점검 업무해설서 요약 Word 문서 생성"""

import sys

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches, Pt, RGBColor


def set_cell_background(cell, fill_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tc_pr.append(shd)


def set_cell_margins(cell, top=90, bottom=90, left=140, right=140) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = OxmlElement("w:tcMar")
    for tag, val in [("w:top", top), ("w:bottom", bottom), ("w:left", left), ("w:right", right)]:
        node = OxmlElement(tag)
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        tc_mar.append(node)
    tc_pr.append(tc_mar)


def add_heading(doc: docx.Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.name = "맑은 고딕"
    run.font.color.rgb = RGBColor(30, 58, 138)
    run.font.size = Pt(14 if level == 1 else 11)


def add_body(doc: docx.Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    run.font.name = "맑은 고딕"
    run.font.size = Pt(9.5)
    run.bold = bold
    run.font.color.rgb = RGBColor(51, 51, 51)


def add_bullet(doc: docx.Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.2
    run = p.add_run(text)
    run.font.name = "맑은 고딕"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(51, 51, 51)


def add_table(doc: docx.Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for col_idx, header in enumerate(headers):
        cell = table.rows[0].cells[col_idx]
        cell.text = header
        set_cell_background(cell, "1E3A8A")
        set_cell_margins(cell)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.name = "맑은 고딕"
                run.font.size = Pt(8.5)
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for row_idx, row_data in enumerate(rows, start=1):
        for col_idx, value in enumerate(row_data):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = value
            set_cell_margins(cell)
            if row_idx % 2 == 0:
                set_cell_background(cell, "F8FAFC")
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "맑은 고딕"
                    run.font.size = Pt(8.5)
                    run.font.color.rgb = RGBColor(51, 51, 51)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)


def add_note_box(doc: docx.Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    cell.width = Inches(6.8)
    set_cell_background(cell, "F0F8FF")
    set_cell_margins(cell, top=100, bottom=100, left=180, right=140)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "right", "bottom"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "none")
        tc_borders.append(b)
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "36")
    left.set(qn("w:color"), "007BFF")
    tc_borders.append(left)
    tc_pr.append(tc_borders)

    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.name = "맑은 고딕"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(85, 85, 85)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)


def build_maintenance_section(doc: docx.Document) -> None:
    add_heading(doc, "Ⅰ. 유지보수·관리 — 진행방법 및 주요 작성항목", level=1)

    add_table(
        doc,
        ["구분", "내용"],
        [
            ["법적 근거", "「정보통신공사업법」 제37조의2, 「정보통신설비 유지보수·관리기준」 제8조"],
            ["정의", "건축물 등 정보통신설비의 기능 유지·이용자 편의·안전 확보를 위한 일상적 보수·관리"],
            ["점검 주기", "완공일 기준 반기 1회 이상 (연 2회)"],
            ["점검 대상", "별표 1 점검대상 정보통신설비 34종"],
            ["점검 항목", "외관 / 기능 / 안전 3개 영역"],
            ["핵심 서식", "별지 제2호 — 정보통신설비 유지보수·관리 점검표"],
        ],
    )

    add_heading(doc, "1. 진행 절차 (5단계)", level=2)
    add_body(
        doc,
        "① 자료 구비 → ② 관리방식 결정 → ③ 유지보수·관리자 선·해임 신고 "
        "→ ④ 계획 수립 → ⑤ 점검 실시 및 기록(반기 1회 이상)",
    )

    add_body(doc, "① 자료 구비", bold=True)
    add_bullet(doc, "정보통신설비 준공도면")
    add_bullet(doc, "설치 현황표 (별지 제1호서식)")

    add_body(doc, "② 관리방식 결정", bold=True)
    add_table(
        doc,
        ["방식", "내용"],
        [
            ["자체관리", "관리주체가 유지보수·관리자(인정교육 이수 기술자)를 선임하여 직접 점검"],
            ["위탁관리", "공사업자에 계획 수립·점검 업무 위탁 (위탁업체는 유지보수·관리자 선임)"],
        ],
    )

    add_body(doc, "③ 유지보수·관리자 선·해임 신고", bold=True)
    add_bullet(doc, "선임·해임 시 관할 지자체(시·군·구청장)에 신고")
    add_bullet(doc, "해임 시 30일 이내 신규 선임")

    add_body(doc, "④ 계획 수립 (최초 점검 전, 변경 시 갱신)", bold=True)
    for item in [
        "대상 현황표 (별지 제1호)",
        "점검 절차 (점검종류·주기·일정)",
        "산업재해방지 대책",
        "긴급상황 매뉴얼 (비상연락망 포함)",
        "이상 상황 발생 시 조치 방법",
        "사고 이력 시 재발방지 대책",
    ]:
        add_bullet(doc, item)

    add_body(doc, "⑤ 점검 실시 및 기록", bold=True)
    add_bullet(doc, "외관·기능·안전 상태를 설비별 점검표에 기록")
    add_bullet(doc, "부적합 시 개선·보수·수리·교체 등 관리주체에 요청 가능")

    doc.add_page_break()

    add_heading(doc, "2. 점검표 작성 방법 (별지 제2호)", level=2)
    add_table(
        doc,
        ["번호", "항목", "작성 내용"],
        [
            ["①", "점검자", "유지보수·관리자 성함"],
            ["②", "설치업체", "현재 기준 설치 업체명"],
            ["③", "설치위치", "도면·현장실사 기준 위치"],
            ["④", "점검항목/내용", "외관·기능·안전 세부 항목"],
            ["⑤", "점검결과", "○(적합) / ×(부적합) / -(해당없음)"],
            ["⑥", "비고", "부적합 사유·특이사항"],
        ],
    )

    add_heading(doc, "3. 점검 3대 영역 (설비 공통)", level=2)
    add_table(
        doc,
        ["영역", "대표 점검 내용"],
        [
            ["외관", "오염·부식·손상·파손, 케이블·커넥터 상태, 고정·취부, 작동표시부(LCD/LED)"],
            ["기능", "정상 동작, 충전량·팬·알람, 통신·수신 상태 등 설비별 기능"],
            ["안전", "설치환경(먼지·습도·온도), 전원·접지, 이상발열·소음, 접지저항 측정, 조명·항온항습"],
        ],
    )
    add_body(doc, "※ 현장 여건에 따라 점검 항목 추가·변경 가능")

    add_heading(doc, "4. 점검 결과서 구성 (부록 예시)", level=2)
    add_table(
        doc,
        ["순서", "구성", "필수"],
        [
            ["1", "유지보수·관리 점검 계획", "●"],
            ["2", "인력 투입 계획 및 장비 현황", ""],
            ["3", "현장 개요 및 관리주체", ""],
            ["4", "설비별 유지보수·관리 점검표", "●"],
            ["5", "점검결과 내역서", ""],
        ],
    )

    add_heading(doc, "5. 핵심 체크리스트", level=2)
    for item in [
        "준공도면·현황표(별지1) 구비",
        "유지보수·관리자 선임 및 지자체 신고",
        "점검 계획서 작성 (5~6개 항목 포함)",
        "반기 1회 설비별 점검표(별지2) 작성",
        "외관·기능·안전 3영역 점검",
        "부적합 → 비고·내역서 기록 → 관리주체 조치 요청",
        "긴급상황·이상상황 매뉴얼·비상연락망 유지",
    ]:
        add_bullet(doc, f"☐ {item}")


def build_performance_section(doc: docx.Document) -> None:
    doc.add_page_break()
    add_heading(doc, "Ⅱ. 성능점검 — 진행방법 및 주요 작성항목", level=1)

    add_table(
        doc,
        ["구분", "내용"],
        [
            ["법적 근거", "「정보통신공사업법」 제37조의3, 「정보통신설비 유지보수·관리기준」 제10조"],
            ["정의", "정보통신설비의 운전·운용에 필요한 성능을 점검하는 것"],
            ["점검 주기", "완공일 기준 매년 1회 이상"],
            ["보존 기간", "성능점검표 5년"],
            ["점검 대상", "유지보수·관리와 동일 34종 설비"],
            ["핵심 서식", "별지 제3호 — 성능점검표 + 별표 2 — 성능점검 검토사항"],
        ],
    )

    add_heading(doc, "1. 진행 절차 (6단계)", level=2)
    add_body(
        doc,
        "① 자료 구비 → ② 점검방식 결정 → ③ 계획 수립 "
        "→ ④ 성능점검 실시(연 1회↑) → ⑤ 기록·보존(5년) → ⑥ 제출(요청 시)",
    )

    add_body(doc, "① 자료 구비", bold=True)
    add_bullet(doc, "준공도면, 설치 현황표(별지1)")
    add_bullet(doc, "성능점검 시 검토사항(별표2)")

    add_body(doc, "② 점검방식 결정", bold=True)
    add_table(
        doc,
        ["방식", "내용"],
        [
            ["자체점검", "연면적에 맞는 등급의 정보통신기술자 고용 후 직접 점검"],
            ["대행점검", "공사업자 또는 용역업자에 대행 위탁"],
        ],
    )

    add_body(doc, "③~⑥ 주요 사항", bold=True)
    add_bullet(doc, "계획: 현황표·절차·재해방지·긴급매뉴얼·이상조치 등 포함")
    add_bullet(doc, "실시: 구비 자료·현황표·검토사항(별표2) 참고, 측정·시험 중심 점검")
    add_bullet(doc, "기록: 별지3 성능점검표 + 별표2 검토사항 작성, 5년 보존")
    add_bullet(doc, "제출: 시·군·구청장 요청 시 성능점검표(전자문서 포함) 제출")

    doc.add_page_break()

    add_heading(doc, "2. 성능점검표 작성 방법 (별지 제3호)", level=2)
    add_table(
        doc,
        ["번호", "항목", "작성 내용"],
        [
            ["①", "점검자(소속)", "성능점검자 또는 대행업체"],
            ["②", "관리주체(입회자)", "입회자 성함"],
            ["③", "설치위치", "도면·현장실사 기준"],
            ["④", "점검항목/내용", "설비별 성능 점검 세부 항목"],
            ["⑤", "점검결과", "○ / × / -"],
            ["⑥", "부적합 항목", "부적합사항 + 조치사항 (× 시 필수)"],
            ["⑦", "현황 사진", "점검대상 설비 사진 첨부"],
            ["⑧", "비고", "특이사항"],
        ],
    )

    add_heading(doc, "3. 유지보수·관리 vs 성능점검 비교", level=2)
    add_table(
        doc,
        ["구분", "유지보수·관리", "성능점검"],
        [
            ["목적", "일상적 상태 유지", "운전·운용 성능 확인"],
            ["주기", "반기 1회↑", "연 1회↑"],
            ["점검 성격", "외관·기능·안전 상태", "측정·시험·설정 확인"],
            ["결과 기록", "점검표(별지2)", "점검표(별지3) + 검토사항(별표2)"],
            ["사진", "선택", "현황 사진 첨부(⑦)"],
            ["보존", "-", "5년"],
        ],
    )

    add_heading(doc, "4. 성능점검 검토사항 (별표 2)", level=2)
    add_body(doc, "① 정보통신설비 시스템 검토", bold=True)
    add_bullet(doc, "34종 설비별 정상 작동 여부 확인·기록")
    add_bullet(doc, "현황표 제조사·모델번호와 현장 설비 일치 여부 대조 (○/×/―)")

    add_body(doc, "② 성능개선 계획 수립", bold=True)
    add_bullet(doc, "내구연수 대비 사용연수(노후도) 분석")
    add_bullet(doc, "점검표 부적합 항목별 개선사항 제시")
    add_bullet(doc, "성능개선 필요성 및 연도별 세부개선계획 수립")
    add_body(doc, "※ 즉시 개선 완료 시 '적합'으로 변경 가능")

    add_heading(doc, "5. 성능점검 결과서 구성 및 체크리스트", level=2)
    add_table(
        doc,
        ["순서", "구성", "필수"],
        [
            ["1", "성능점검 계획", "●"],
            ["2", "인력·장비 현황", ""],
            ["3", "현장 개요·관리주체", ""],
            ["4", "설비별 성능점검표(별지3)", "●"],
            ["5", "성능점검 검토사항(별표2)", "●"],
            ["6", "점검결과 내역서", ""],
            ["7", "성능개선 계획", ""],
        ],
    )

    for item in [
        "준공도면·현황표·검토사항(별표2) 구비",
        "자체/대행 방식 결정 (자체 시 자격 등급 확인)",
        "연 1회 34종 설비별 성능점검표(별지3) 작성",
        "× 항목 → 부적합사항 + 조치사항 상세 기재",
        "현황 사진 첨부",
        "별표2 시스템 검토 + 성능개선 계획 작성",
        "결과서 5년 보존, 지자체 요청 시 제출",
    ]:
        add_bullet(doc, f"☐ {item}")


def build_document(output_path: str) -> None:
    doc = docx.Document()

    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.font.size = Pt(9.5)
    style.font.color.rgb = RGBColor(51, 51, 51)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(6)
    title_run = title_p.add_run("정보통신설비 유지보수·관리 및 성능점검\n업무 요약서")
    title_run.font.name = "맑은 고딕"
    title_run.font.size = Pt(16)
    title_run.bold = True
    title_run.font.color.rgb = RGBColor(30, 58, 138)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_after = Pt(10)
    sub_run = sub_p.add_run("근거: 정보통신설비 유지보수·관리 및 성능점검 업무 해설서 (2025. 8.)")
    sub_run.font.name = "맑은 고딕"
    sub_run.font.size = Pt(9)
    sub_run.font.color.rgb = RGBColor(100, 100, 100)

    add_note_box(
        doc,
        "본 문서는 업무 해설서를 바탕으로 작성된 참고용 요약 자료입니다. "
        "법적 근거로 사용될 수 없으며, 실제 작성 시 「정보통신설비 유지보수·관리기준」 "
        "별지·별표 서식을 반드시 함께 확인하시기 바랍니다.",
    )

    build_maintenance_section(doc)
    build_performance_section(doc)

    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_p.paragraph_format.space_before = Pt(12)
    footer_run = footer_p.add_run("작성: Gro | 2026-06-12")
    footer_run.font.name = "맑은 고딕"
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(150, 150, 150)

    doc.save(output_path)


def main() -> None:
    sys.stdout.reconfigure(encoding="utf-8")
    output_path = r"D:\Dev\Project\sample\정보통신설비_유지보수·성능점검_업무요약서.docx"
    build_document(output_path)
    print(f"Word 문서 생성 완료: {output_path}")


if __name__ == "__main__":
    main()
