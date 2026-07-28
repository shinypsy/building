import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    """표 셀 배경색 채우기"""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=140, bottom=140, left=180, right=180):
    """표 셀 여백(Padding) 설정 (dxa 단위, 1pt = 20 dxa)"""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_borders(cell, **kwargs):
    """셀 개별 테두리 설정 (left, top, right, bottom)"""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        edge_data = kwargs.get(edge)
        if edge_data:
            b = OxmlElement(f'w:{edge}')
            b.set(qn('w:val'), edge_data.get('val', 'single'))
            b.set(qn('w:sz'), str(edge_data.get('sz', 4)))
            b.set(qn('w:space'), '0')
            b.set(qn('w:color'), edge_data.get('color', 'auto'))
            tcBorders.append(b)
        else:
            b = OxmlElement(f'w:{edge}')
            b.set(qn('w:val'), 'none')
            tcBorders.append(b)
    tcPr.append(tcBorders)

def build_fluke_prices_docx():
    doc = docx.Document()
    
    # 상하좌우 여백 1.0인치 설정
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # 기본 스타일 정의 (맑은 고딕)
    style = doc.styles['Normal']
    font = style.font
    font.name = '맑은 고딕'
    font.size = Pt(10)
    font.color.rgb = RGBColor(51, 51, 51) # Charcoal
    
    # 1. 제목 (Title)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    run = p.add_run("Fluke Networks 케이블 측정기 U.S. 실거래가 및 견적 조사 보고서")
    run.font.size = Pt(16)
    run.bold = True
    run.font.color.rgb = RGBColor(30, 58, 138) # Deep Navy (1E3A8A)
    
    # 2. 리서치 배경 및 개요
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.3
    run_bold = p.add_run("1. 리서치 배경 및 주요 유통 경로\n")
    run_bold.bold = True
    run_bold.font.size = Pt(11)
    run_bold.font.color.rgb = RGBColor(30, 58, 138)
    
    p.add_run(
        "본 보고서는 미국 내 주요 공식 Authorized Distributor(TEquipment, Global Test Supply, ITM.com, "
        "Test Equipment Depot 등)의 실시간 판매 단가 및 납품 견적 자료를 모델별로 조사하여 작성되었습니다.\n"
        "Fluke Networks의 케이블 측정 장비는 네트워크 구축 품질 인증 및 유지보수에 핵심적인 장비로, "
        "미국 내에서도 장비 고유의 정밀도와 신뢰성으로 높은 가치를 형성하고 있습니다. 조사된 실거래 가격은 "
        "구성 패키지(Bundle), WiFi 모듈 여부, 사후 지원 서비스(Gold Support) 등에 따라 다르게 책정될 수 있습니다."
    )
    
    # 3. 가격 상세 분석 (표 도입)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(10)
    run_bold = p.add_run("2. 주요 모델별 미국 내 가격 현황\n")
    run_bold.bold = True
    run_bold.font.size = Pt(11)
    run_bold.font.color.rgb = RGBColor(30, 58, 138)
    
    # 표 생성 (6행 4열: 헤더 + 5개 제품)
    table = doc.add_table(rows=6, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # 열 너비 설정
    widths = [Inches(2.5), Inches(1.3), Inches(1.3), Inches(1.4)]
    
    headers = ["모델명 및 주요 사양", "공식 정가 (List)", "실거래가 (Discounted)", "중고/리퍼브 시세"]
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], "1E3A8A") # Deep Navy
        set_cell_margins(hdr_cells[i], top=140, bottom=140, left=100, right=100)
        
        # 글자 스타일 설정
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.font.name = '맑은 고딕'
        run.font.size = Pt(9.5)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        
        # 헤더 테두리
        set_cell_borders(hdr_cells[i], 
                         top={"val": "single", "sz": 8, "color": "1E3A8A"},
                         bottom={"val": "single", "sz": 12, "color": "1E3A8A"},
                         left={"val": "none"}, right={"val": "none"})
        
    data = [
        (
            "DSX2-8000 CableAnalyzer\n(Versiv 2 기반 Cat 8 인증용 구리선 테스터)", 
            "약 $16,080", 
            "약 $15,063", 
            "약 $11,000 ~ $15,000"
        ),
        (
            "DSX2-5000 CableAnalyzer\n(Versiv 2 기반 Cat 6A 인증용 구리선 테스터)", 
            "약 $14,225\n(Non-Wireless)", 
            "약 $13,325\n(Wi-Fi 탑재형)", 
            "약 $6,500 ~ $8,500"
        ),
        (
            "LinkIQ Cable + Network Tester\n(10 Gb/s 대역폭 검증 및 PoE 실측 테스터)", 
            "약 $2,661\n(LIQ-100 단품)", 
            "약 $2,200 ~ $2,661\n(키트 $3,100~$3,600)", 
            "약 $1,900 ~ $3,000"
        ),
        (
            "MicroScanner2 Series\n(Wiremap, PoE 검사용 보급형 케이블 테스터)", 
            "약 $790\n(MS2-100 단품)", 
            "약 $680 ~ $790\n(키트 $900~$1,350)", 
            "약 $350 ~ $600"
        ),
        (
            "CertiFiber Pro (CFP2-100-S)\n(싱글모드 광케이블 손실 측정기 키트)", 
            "약 $12,500\n(예상 정가)", 
            "약 $11,398\n(Quad키트 $28,000~)", 
            "약 $8,000 ~ $10,500"
        )
    ]
    
    border_color = "D1D5DB" # Light Gray
    for row_idx, row_data in enumerate(data, start=1):
        row_cells = table.rows[row_idx].cells
        
        # 지브라 패턴 (짝수 행 배경색 삽입)
        bg_color = "F3F4F6" if row_idx % 2 == 0 else "FFFFFF"
        
        for col_idx, text in enumerate(row_data):
            row_cells[col_idx].text = text
            set_cell_background(row_cells[col_idx], bg_color)
            set_cell_margins(row_cells[col_idx], top=120, bottom=120, left=100, right=100)
            
            p = row_cells[col_idx].paragraphs[0]
            if col_idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
            for run in p.runs:
                run.font.name = '맑은 고딕'
                run.font.size = Pt(9)
                
            # 셀 테두리 설정 (아래쪽 얇은 테두리)
            set_cell_borders(row_cells[col_idx],
                             bottom={"val": "single", "sz": 4, "color": border_color},
                             top={"val": "none"}, left={"val": "none"}, right={"val": "none"})
            
    # 전체 열 너비 맞춤
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width

    # 4. 강조 Callout Box (구매 및 견적 시 주요 검토사항)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(6)
    run_bold = p.add_run("3. 구매 및 견적 시 주요 검토사항")
    run_bold.bold = True
    run_bold.font.size = Pt(11)
    run_bold.font.color.rgb = RGBColor(30, 58, 138)

    # 1x1 테이블을 사용하여 Callout Box 생성
    callout_table = doc.add_table(rows=1, cols=1)
    callout_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    callout_cell = callout_table.rows[0].cells[0]
    callout_cell.width = Inches(6.5)
    
    set_cell_background(callout_cell, "EBF3FC") # 옅은 청색 배경
    set_cell_margins(callout_cell, top=160, bottom=160, left=200, right=200)
    
    # 좌측에 굵은 테두리(sz=24 -> 3pt), 다른 면은 없음
    set_cell_borders(callout_cell, 
                     left={"val": "single", "sz": 24, "color": "1E3A8A"},
                     top={"val": "none"}, right={"val": "none"}, bottom={"val": "none"})
                     
    callout_p = callout_cell.paragraphs[0]
    callout_p.paragraph_format.line_spacing = 1.3
    
    run1 = callout_p.add_run("💡 성공적인 도입을 위한 핵심 체크포인트\n")
    run1.bold = True
    run1.font.size = Pt(9.5)
    run1.font.color.rgb = RGBColor(30, 58, 138)
    
    bullets = [
        "Gold Support 계약 여부: Fluke Networks의 프리미엄 사후 관리 플랜(연간 교정, 부품 무상 교체 등) 포함 여부에 따라 대당 약 $1,000 이상의 견적 차이가 발생하므로 기안 시 Gold Support 필요 항목을 필히 정의해야 합니다.",
        "기관 및 대량 구매 할인(EDU/GOV): 미국 디스트리뷰터들은 학교, 관공서, 군납 또는 일정 수량 이상 구매 시 추가 5%~10% 특별 할인 견적을 제공하므로 바잉 파워(Buying Power)를 활용한 특별 견적 요청이 유리합니다.",
        "정밀 부가 모듈 구성 검토: 단순 본체 구매 외에 WiFi 동글, 스마트 원격 수신기, 광 단면 검사용 프로브(FI-7000 등) 등의 액세서리 세트가 포함되었는지 확인하고 최적의 일괄 계약을 진행해야 합니다."
    ]
    
    for b in bullets:
        run_bullet = callout_p.add_run(f"• {b}\n")
        run_bullet.font.size = Pt(9)
        run_bullet.font.color.rgb = RGBColor(51, 51, 51)
        
    if callout_p.runs:
        callout_p.runs[-1].text = callout_p.runs[-1].text.rstrip('\n')
        
    doc.save("Fluke_케이블측정기_미국실거래가_조사보고서.docx")
    print("보고서 워드 파일(Fluke_케이블측정기_미국실거래가_조사보고서.docx) 생성 완료!")

if __name__ == '__main__':
    build_fluke_prices_docx()
