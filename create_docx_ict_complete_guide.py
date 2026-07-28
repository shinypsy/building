"""정보통신설비 유지보수·관리 완전 가이드 Word 문서 생성"""

import sys
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches, Pt, RGBColor


def set_cell_background(cell, fill_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tc_pr.append(shd)


def set_cell_margins(cell, top=90, bottom=90, left=140, right=140) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = OxmlElement("w:tcMar")
    for tag, val in [("w:top", top), ("w:bottom", bottom), ("w:left", left), ("w:right", right)]:
        node = OxmlElement(tag)
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        tc_mar.append(node)
    tc_pr.append(tc_mar)


def add_heading(doc: docx.Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.name = "맑은 고딕"
    run.font.color.rgb = RGBColor(30, 58, 138)
    run.font.size = Pt(14 if level == 1 else 11)


def add_body(doc: docx.Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    run.font.name = "맑은 고딕"
    run.font.size = Pt(9.5)
    run.bold = bold
    run.font.color.rgb = RGBColor(51, 51, 51)


def add_bullet(doc: docx.Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.2
    run = p.add_run(text)
    run.font.name = "맑은 고딕"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(51, 51, 51)


def add_table(doc: docx.Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for col_idx, header in enumerate(headers):
        cell = table.rows[0].cells[col_idx]
        cell.text = header
        set_cell_background(cell, "1E3A8A")
        set_cell_margins(cell)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.name = "맑은 고딕"
                run.font.size = Pt(8.5)
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for row_idx, row_data in enumerate(rows, start=1):
        for col_idx, value in enumerate(row_data):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = value
            set_cell_margins(cell)
            if row_idx % 2 == 0:
                set_cell_background(cell, "F8FAFC")
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "맑은 고딕"
                    run.font.size = Pt(8.5)
                    run.font.color.rgb = RGBColor(51, 51, 51)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)


def build_document(output_path: str) -> None:
    doc = docx.Document()

    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.font.size = Pt(9.5)
    style.font.color.rgb = RGBColor(51, 51, 51)

    # 제목
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(6)
    title_run = title_p.add_run("CCTV에서 스마트공장까지\n정보통신설비 유지보수 대상 34종과 점검 절차 총정리")
    title_run.font.name = "맑은 고딕"
    title_run.font.size = Pt(16)
    title_run.bold = True
    title_run.font.color.rgb = RGBColor(30, 58, 138)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_after = Pt(10)
    sub_run = sub_p.add_run("원문: https://jackti.tistory.com/2255")
    sub_run.font.name = "맑은 고딕"
    sub_run.font.size = Pt(9)
    sub_run.font.color.rgb = RGBColor(100, 100, 100)

    # PART 1
    add_heading(doc, "PART 1. 제도의 전체 구조", level=1)
    add_heading(doc, "1. 정보통신설비 유지보수·관리 제도란?", level=2)

    add_body(doc, "1-1. 제도의 목적", bold=True)
    add_bullet(doc, "건축물·시설물 등에 설치된 정보통신설비의 고장 및 훼손, 방치 문제를 예방")
    add_bullet(doc, "관리주체가 정보통신설비를 유지보수·관리하고 성능을 점검")
    add_bullet(doc, "국민의 안전과 정보통신서비스의 원활한 제공이 목표")

    add_body(doc, "1-2. 대상 건축물", bold=True)
    add_table(
        doc,
        ["구분", "내용"],
        [
            ["대상", "연면적 5,000㎡ 이상의 건축물"],
            ["제외", "공동주택"],
            ["제외", "학교시설"],
        ],
    )

    add_body(doc, "1-3. 두 가지 의무", bold=True)
    add_table(
        doc,
        ["구분", "유지보수·관리", "성능점검"],
        [
            ["목적", "일상적 보수·관리", "운용 성능 확인"],
            ["주기", "반기별 1회 이상", "매년 1회 이상"],
            ["기록", "점검표 작성", "성능점검표 작성 및 5년 보존"],
        ],
    )

    add_body(doc, "핵심 차이: 유지보수·관리는 '설비가 고장 없이 작동하는가?'를 점검하고, 성능점검은 '운용 성능이 기준에 부합하는가?'를 평가합니다.")

    doc.add_page_break()

    # PART 2
    add_heading(doc, "PART 2. 유지보수 대상 정보통신설비 34종", level=1)
    add_heading(doc, "2. 34종 설비 분류", level=2)

    add_table(
        doc,
        ["분류", "종수", "대표 설비"],
        [
            ["통신설비", "8종", "케이블, 배관, 전화, 이동통신 등"],
            ["방송설비", "1종", "방송음향설비"],
            ["정보설비", "23종", "CCTV, 주차관제, BEMS, 스마트공장 등"],
            ["기타설비", "2종", "통신용 전원, 통신접지"],
        ],
    )

    add_heading(doc, "2-1. 통신설비 (8종)", level=2)
    add_body(doc, "건축물 내외부의 음성·데이터 통신을 위한 기반 인프라. 건물의 '신경망'에 해당합니다.")

    add_table(
        doc,
        ["설비명", "점검 포인트", "흔한 문제"],
        [
            ["케이블설비", "케이블 손상, 접속 상태", "노후화로 인한 신호 감쇠"],
            ["배관설비", "파손, 이음새, 습기 유입", "균열로 인한 침수"],
            ["국선인입설비", "인입구 보호, 방수 처리", "노후화로 인한 외부 침수"],
            ["단자함설비", "내부 접속 상태, 습기", "단자 부식, 접속 불량"],
            ["이동통신구내선로", "중계기 작동, 안테나", "중계기 오작동"],
            ["전화설비", "교환기, 전화기, 배선", "교환기 노후화"],
            ["방송공동수신안테나", "안테나 설치, 증폭기", "안테나 방향 틀어짐"],
            ["종합유선방송구내전송", "전송선로, 분배기", "선로 노후화"],
        ],
    )

    add_heading(doc, "2-2. 방송설비 (1종)", level=2)
    add_body(doc, "건축물 내부의 음향 방송 시스템. 비상방송 기능은 화재 시 대피 안내에 직접 관련되므로 매우 중요합니다.")

    add_body(doc, "방송음향설비", bold=True)
    add_bullet(doc, "점검: 스피커 작동, 앰프 출력, 마이크 음질, 비상방송 기능")
    add_bullet(doc, "문제: 스피커 고장, 앰프 과부하, 비상방송 미작동")

    add_heading(doc, "2-3. 정보설비 (23종)", level=2)
    add_body(doc, "건물의 지능화·자동화를 담당하는 핵심 설비입니다.")

    add_body(doc, "보안·출입 분야", bold=True)
    add_bullet(doc, "네트워크설비: 스위치·허브·무선AP 작동 상태")
    add_bullet(doc, "전자출입시스템: 리더기 인식률, 도어락 작동")
    add_bullet(doc, "CCTV: 화질, 녹화 상태, 저장장치, 야간 촬영")
    add_bullet(doc, "비상벨설비: 벨 작동, 음량, 연결 상태")
    add_bullet(doc, "지능형 경계감시: 센서 감도, AI 분석 정확도")

    add_body(doc, "주차·물류 분야", bold=True)
    add_bullet(doc, "주차관제시스템: 차량 번호 인식률, 게이트 작동")
    add_bullet(doc, "주차유도시스템: 센서 감지 정확도, 안내판 표시")
    add_bullet(doc, "무인택배시스템: 도어 작동, 터치스크린, 서버 연결")

    add_body(doc, "건물관리·에너지 분야", bold=True)
    add_bullet(doc, "시설관리시스템(FMS): 각 설비와의 연동, 데이터 수집")
    add_bullet(doc, "건물에너지관리시스템(BEMS): 센서 정확도, 제어 기능")
    add_bullet(doc, "원격검침시스템: 검침 센서 정확도, 통신 상태")
    add_bullet(doc, "빌딩안내시스템(BIS): 디스플레이 작동, 콘텐츠 업데이트")

    add_body(doc, "스마트 기술 분야", bold=True)
    add_bullet(doc, "홈네트워크 설비: 월패드, IoT 기기 연동")
    add_bullet(doc, "지능형 인원계수: 센서 인식 정확도")
    add_bullet(doc, "스마트 병원설비(너스콜): 호출 버튼, 통화 품질")
    add_bullet(doc, "스마트 공장: IoT 센서, 데이터 수집·분석")
    add_bullet(doc, "지능형 이상음원: 마이크 감도, AI 분석 정확도")
    add_bullet(doc, "IoT 지하공간 안전관리: 가스·수위·화재 센서")

    add_heading(doc, "2-4. 기타설비 (2종)", level=2)
    add_body(doc, "통신의 '심장' 역할을 하는 설비입니다. 전원이 끊기면 모든 정보통신설비가 정지합니다.")

    add_table(
        doc,
        ["설비명", "점검 포인트", "중요성"],
        [
            ["통신용 전원설비", "UPS 배터리, 정류기 출력", "가장 기초적인 점검 대상"],
            ["통신접지설비", "접지 저항값, 접지선 연결", "고장과 안전사고의 주요 원인"],
        ],
    )

    doc.add_page_break()

    # PART 3
    add_heading(doc, "PART 3. 점검 절차", level=1)
    add_heading(doc, "3. 유지보수·관리 점검 절차", level=2)

    add_body(doc, "점검 5단계", bold=True)
    add_bullet(doc, "① 자료 구비 (준공도면, 설치 현황표)")
    add_bullet(doc, "② 관리방식 결정 (자체관리 vs 위탁관리)")
    add_bullet(doc, "③ 유지보수·관리자 선·해임 신고 (지자체 신고)")
    add_bullet(doc, "④ 계획 수립 (점검 절차, 재해방지 대책, 긴급 매뉴얼)")
    add_bullet(doc, "⑤ 점검 실시 및 기록 (반기 1회 이상)")

    add_body(doc, "점검 3대 영역", bold=True)
    add_table(
        doc,
        ["영역", "대표 점검 내용"],
        [
            ["외관", "오염·부식·손상·파손, 케이블·커넥터 상태, 고정·취부"],
            ["기능", "정상 동작, 충전량·팬·알람, 통신·수신 상태"],
            ["안전", "설치환경, 전원·접지, 이상발열·소음, 조명·항온항습"],
        ],
    )

    add_heading(doc, "3-1. 성능점검 절차", level=2)

    add_body(doc, "성능점검 6단계", bold=True)
    add_bullet(doc, "① 성능점검 계획서 작성")
    add_bullet(doc, "② 성능점검 실시 (연 1회 이상)")
    add_bullet(doc, "③ 성능점검표 기록")
    add_bullet(doc, "④ 성능개선 계획 수립")
    add_bullet(doc, "⑤ 기록 5년간 보존")
    add_bullet(doc, "⑥ 지자체 요청 시 제출")

    add_body(doc, "성능개선 계획의 주요 항목", bold=True)
    add_bullet(doc, "내구연수 대비 사용연수(노후도) 분석")
    add_bullet(doc, "점검표 부적합 항목별 개선사항")
    add_bullet(doc, "연도별 세부 개선 계획")

    doc.add_page_break()

    # PART 4
    add_heading(doc, "PART 4. 계획서와 관리자 자격", level=1)
    add_heading(doc, "4. 계획서에 포함할 8가지 필수 항목", level=2)

    add_table(
        doc,
        ["번호", "항목"],
        [
            ["1", "점검대상 정보통신설비의 종류 및 항목"],
            ["2", "유지보수·관리 및 성능점검 절차와 주기"],
            ["3", "유지보수·관리 및 성능점검 전 재해방지 대책"],
            ["4", "긴급 상황에 대한 매뉴얼"],
            ["5", "사고·이상 상황 발생 시 조치 방법 및 재발방지 대책"],
            ["6", "성능점검을 위한 인력 투입 계획 및 장비 현황"],
            ["7", "안전 확보 및 품질 관리 방안"],
            ["8", "성능점검 결과에 따른 조치 방안"],
        ],
    )

    add_heading(doc, "4-1. 관리자 자격기준 (규모별)", level=2)

    add_table(
        doc,
        ["건축물 규모 (연면적)", "적용 시점", "관리자 자격"],
        [
            ["60,000㎡ 이상", "2025. 7. 19.", "특급 기술자"],
            ["30,000㎡ ~ 60,000㎡ 미만", "2025. 7. 19.", "고급 기술자 이상"],
            ["15,000㎡ ~ 30,000㎡ 미만", "2026. 7. 19.", "중급 기술자 이상"],
            ["10,000㎡ ~ 15,000㎡ 미만", "2026. 7. 19.", "초급 기술자 이상"],
            ["5,000㎡ ~ 10,000㎡ 미만", "2027. 7. 19.", "초급 기술자 이상"],
        ],
    )

    add_body(doc, "인정교육 필수: 기술계 정보통신기술자 자격 취득 후 20시간 이상 인정교육 이수 필요")
    add_body(doc, "중복 선임 가능: 1명의 관리자가 최대 5개의 건축물에 중복 선임 가능")

    doc.add_page_break()

    # PART 5
    add_heading(doc, "PART 5. 업무 위탁과 과태료", level=1)
    add_heading(doc, "5-1. 업무 위탁", level=2)

    add_table(
        doc,
        ["구분", "위탁 대상"],
        [
            ["유지보수·관리", "정보통신공사업자"],
            ["성능점검", "정보통신공사업자, 통신관련 엔지니어링 사업자, 기술사사무소"],
        ],
    )

    add_body(doc, "⚠️ 위탁 시 관리주체의 잔여 의무", bold=True)
    add_bullet(doc, "계획서 작성 의무는 여전히 유지")
    add_bullet(doc, "현황표 작성·비치·현행화 의무는 여전히 유지")
    add_bullet(doc, "위탁은 관리자 선임 의무만 면제")

    add_heading(doc, "5-2. 과태료 체계", level=2)

    add_table(
        doc,
        ["위반행위", "과태료"],
        [
            ["유지보수·관리기준 미준수", "300만 원"],
            ["점검기록 미작성·거짓 작성", "300만 원"],
            ["유지보수·관리자 미선임", "300만 원"],
            ["관리자 해임 후 30일 이내 후임 미선임", "300만 원"],
            ["점검기록 미보존", "150만 원"],
            ["점검기록 지자체 미제출", "100만 원"],
            ["선·해임 신고 미신고·거짓 신고", "100만 원"],
        ],
    )

    add_body(doc, "📅 과태료 유예", bold=True)
    add_bullet(doc, "2026년 7월 18일까지: 시정명령·행정지도 등 개선권고")
    add_bullet(doc, "2026년 1월 18일까지: 관리자 선임 및 1회차 유지관리점검 필수")
    add_bullet(doc, "2026년 7월 18일까지: 성능점검 1회, 2회차 유지관리점검 필수")

    doc.add_page_break()

    # PART 6
    add_heading(doc, "PART 6. 선·해임 신고 절차", level=1)

    add_body(doc, "신고 기한: 선임·해임 시 30일 이내에 시·군·구청에 신고", bold=True)

    add_heading(doc, "6-1. 선임 신고 시 제출서류", level=2)

    add_table(
        doc,
        ["번호", "제출서류"],
        [
            ["1", "유지보수·관리자 선임·해임 신고서"],
            ["2", "유지보수·관리자 재직증명서 (직접선임 시) 또는 위탁업무계약서 사본"],
            ["3", "정보통신기술자 경력수첩 사본 및 경력확인서"],
            ["4", "유지보수·관리자 인정교육 20시간 이상 수료증"],
            ["5", "사업자등록증 사본 (행정정보공동이용 미동의 시)"],
            ["6", "건축물대장 (행정정보공동이용 미동의 시)"],
            ["7", "위임장 (대리인 신청 시)"],
        ],
    )

    doc.add_page_break()

    # PART 7
    add_heading(doc, "PART 7. 실전 점검 포인트", level=1)
    add_heading(doc, "7-1. 설비 분류별 핵심 점검", level=2)

    add_body(doc, "통신설비: '연결'이 핵심", bold=True)
    add_bullet(doc, "물리적 연결 확인: 케이블 파손, 접속 이완, 단자 부식")
    add_bullet(doc, "신호 품질 확인: 감쇠율, 반사율, 신호 대 잡음비")

    add_body(doc, "방송설비: '안전'이 핵심", bold=True)
    add_bullet(doc, "비상방송 자동 발동 여부")
    add_bullet(doc, "모든 구역으로의 전달 여부")

    add_body(doc, "정보설비: '정확성'이 핵심", bold=True)
    add_bullet(doc, "데이터 정확성 확인")
    add_bullet(doc, "네트워크 연결 상태 확인")

    add_body(doc, "기타설비: '기본'이 핵심", bold=True)
    add_bullet(doc, "UPS 배터리 수명")
    add_bullet(doc, "접지 저항값 및 접지선 상태")

    add_heading(doc, "7-2. 점검 빈도 권장사항", level=2)

    add_table(
        doc,
        ["점검 빈도", "대상 설비", "이유"],
        [
            ["월 1회 이상", "CCTV, 비상벨, 전자출입, 전원설비", "안전과 직결"],
            ["반기 1회", "대부분의 설비", "법정 최소 기준"],
            ["연 1회", "34종 전체", "성능 종합 평가"],
            ["이벤트 발생 시", "IoT 센서, 이상음원", "이상 감지 시 즉시"],
        ],
    )

    doc.add_page_break()

    # PART 8
    add_heading(doc, "PART 8. 자주 묻는 질문 (FAQ)", level=1)

    add_body(doc, "Q1. 34종 설비 중 없는 설비도 점검해야 하나?", bold=True)
    add_body(doc, "A. 아니요. 점검 대상은 실제 설치된 설비입니다. 없는 설비는 현황표에서 명확히 기록합니다.")

    add_body(doc, "Q2. 유지보수·관리와 성능점검을 동시에 할 수 있나?", bold=True)
    add_body(doc, "A. 일부 겹칠 수 있지만, 법적으로 두 제도는 별개입니다. 성능점검은 별도로 연 1회 이상 실시해야 합니다.")

    add_body(doc, "Q3. 위탁업체가 점검하면 관리주체는 확인할 필요가 없나?", bold=True)
    add_body(doc, "A. 아니요. 위탁하더라도 관리주체는 결과를 확인하고 보수·조치를 이행해야 합니다.")

    add_body(doc, "Q4. 성능점검표를 분실했으면?", bold=True)
    add_body(doc, "A. 5년 보존 의무이므로 기록을 복구해야 합니다. 위탁업체에 보존 여부를 확인하세요.")

    add_body(doc, "Q5. 점검 중 고장을 발견하면 즉시 보수해야 하나?", bold=True)
    add_body(doc, "A. 예, 특히 안전 관련 설비(비상벨, CCTV, 전자출입)는 즉시 조치가 필요합니다.")

    add_body(doc, "Q6. 점검 내용을 추가·변경할 수 있나?", bold=True)
    add_body(doc, "A. 예. 건축물의 설비 특성에 맞게 점검 항목을 맞춤형으로 조정할 수 있습니다.")

    doc.add_page_break()

    # PART 9
    add_heading(doc, "PART 9. 종합 정리", level=1)
    add_heading(doc, "8-1. 34종 설비 분류별 요약", level=2)

    add_table(
        doc,
        ["분류", "종수", "핵심 점검"],
        [
            ["통신설비", "8종", "연결 상태, 신호 품질"],
            ["방송설비", "1종", "비상방송 기능 (안전 중심)"],
            ["정보설비", "23종", "데이터 정확성, 네트워크 연결"],
            ["기타설비", "2종", "전원, 접지 (기반 인프라)"],
        ],
    )

    add_heading(doc, "8-2. 점검 절차 한눈에 보기", level=2)

    add_table(
        doc,
        ["구분", "유지보수·관리", "성능점검"],
        [
            ["주기", "반기별 1회 이상", "연 1회 이상"],
            ["내용", "외관·기능·안전", "운전·운용 성능"],
            ["기록", "점검표 작성", "성능점검표 작성"],
            ["보존", "3~5년 (권장)", "5년 (의무)"],
            ["위탁", "공사업자", "공사업자·용역업자"],
        ],
    )

    add_heading(doc, "8-3. 관리주체의 핵심 의무 체크리스트", level=2)

    for item in [
        "연초: 유지보수·관리 및 성능점검 계획서 작성",
        "변경 시: 현황표 갱신",
        "30일 이내: 관리자 선임 및 지자체 신고",
        "반기 1회 이상: 유지보수·관리 점검 및 기록",
        "연 1회 이상: 성능점검 실시 및 기록",
        "상시: 성능점검 기록 5년 보존",
        "요청 시: 성능점검표 지자체 제출",
    ]:
        add_bullet(doc, f"☐ {item}")

    add_heading(doc, "8-4. 마무리", level=2)

    add_body(doc, "CCTV부터 스마트공장까지, 건축물 안에 설치된 34종의 정보통신설비는 현대 건물의 '신경망'이자 '감각'입니다. 이 설备들이 제대로 작동하지 않으면 건물의 보안·안전·에너지 효율·사용자 편의가 모두 무너집니다.")

    add_body(doc, "4가지 핵심 정리:", bold=True)
    add_bullet(doc, "첫째, 34종 설비의 존재를 파악하고 현황표를 작성합니다.")
    add_bullet(doc, "둘째, 유지보수·관리(반기별)와 성능점검(연 1회)은 별개의 의무입니다.")
    add_bullet(doc, "셋째, 성능점검 기록은 5년간 보존해야 합니다.")
    add_bullet(doc, "넷째, 연초에 계획서를 세우고 누락을 방지합니다.")

    # 면책조항
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run("면책조항")
    run.bold = True
    run.font.size = Pt(10)

    add_body(doc, "본 문서는 블로그 글을 정리한 참고용 자료입니다. 최신 법령은 반드시 법제처 국가법령정보센터(https://www.law.go.kr) 및 과학기술정보통신부(https://www.msit.go.kr)에서 확인하시기 바랍니다.")

    # 작성 정보
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_p.paragraph_format.space_before = Pt(12)
    footer_run = footer_p.add_run("작성: Gro | 2026-06-15")
    footer_run.font.name = "맑은 고딕"
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(150, 150, 150)

    doc.save(output_path)


def main() -> None:
    sys.stdout.reconfigure(encoding="utf-8")
    output_path = r"D:\Dev\Project\sample\정보통신설비_유지보수관리_완전가이드.docx"
    build_document(output_path)
    print(f"Word 문서 생성 완료: {output_path}")


if __name__ == "__main__":
    main()
